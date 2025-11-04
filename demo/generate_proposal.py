"""
Generate a professional Word document proposal with Sinnextech branding for ERP System
Client: Hisham Traders
Requires: python-docx library (install with: pip install python-docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_footer(doc):
    """Add Sinnextech footer to all pages"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""

    # Company description
    desc_run = footer_para.add_run(
        'Sinnextech delivers IT and branding solutions, including HMS, Pharmacy Apps, Lab Integration, '
        'ERP, Web Design, and Brand Identity for healthcare and enterprise sectors.'
    )
    desc_run.font.size = Pt(8)
    desc_run.font.color.rgb = RGBColor(102, 102, 102)

    footer_para.add_run('\n')

    # Contact info table
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6))
    footer_table.autofit = False

    # Column 1: Empty for spacing
    col1 = footer_table.cell(0, 0)
    col1.text = ""

    # Column 2: Office address
    col2 = footer_table.cell(0, 1)
    col2_para = col2.paragraphs[0]
    office_run = col2_para.add_run('Headoffice:\nSinnexTech Office 347, M. Dubai Tower Islamabad')
    office_run.font.size = Pt(8)
    office_run.font.color.rgb = RGBColor(102, 102, 102)

    # Column 3: Contact
    col3 = footer_table.cell(0, 2)
    col3_para = col3.paragraphs[0]
    contact_run = col3_para.add_run(
        'Contact:\nTel +923499724405\nEmail: sinnextech@gmail.com\nWeb: www.sinnextech.com'
    )
    contact_run.font.size = Pt(8)
    contact_run.font.color.rgb = RGBColor(102, 102, 102)

def create_proposal():
    doc = Document()

    # Set default font - Inter style (using Calibri as fallback)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ==================== PAGE 1: COVER PAGE ====================
    # Add spacing
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(8):
        cover_para.add_run().add_break()

    # Logo text (since we can't add actual logo easily)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_run = logo_para.add_run('SinnexTech')
    logo_run.font.size = Pt(24)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(102, 126, 234)  # Purple

    # Add more spacing
    for _ in range(6):
        cover_para.add_run().add_break()

    # Client name badge
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    client_run = client_para.add_run('Hisham Traders')
    client_run.font.size = Pt(14)
    client_run.font.color.rgb = RGBColor(102, 102, 102)

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run('ERP System\nDevelopment Proposal')
    title_run.font.size = Pt(42)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Underline decoration
    decoration = doc.add_paragraph()
    decoration.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dec_run = decoration.add_run('_' * 30)
    dec_run.font.color.rgb = RGBColor(102, 126, 234)
    dec_run.font.size = Pt(16)

    # Spacing
    for _ in range(4):
        decoration.add_run().add_break()

    # Prepared by
    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.LEFT
    prepared_run = prepared.add_run('Prepared By Sinnextech')
    prepared_run.font.size = Pt(16)
    prepared_run.font.color.rgb = RGBColor(0, 0, 0)

    add_page_break(doc)

    # ==================== PAGE 2: PROJECT OVERVIEW ====================
    # Logo header for internal pages
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    # Section heading with purple left border effect (using formatting)
    heading1 = doc.add_heading('ERP System\nDevelopment Proposal', level=1)
    heading1_run = heading1.runs[0]
    heading1_run.font.size = Pt(28)
    heading1_run.font.bold = True
    heading1_run.font.color.rgb = RGBColor(0, 0, 0)

    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run('Prepared for: ').bold = True
    meta_para.add_run('Hisham Traders\n')
    meta_para.add_run('Prepared by: ').bold = True
    meta_para.add_run('Sinnextech\n')
    meta_para.add_run('Date: ').bold = True
    meta_para.add_run('January 2025')

    doc.add_paragraph()

    # Project Overview
    doc.add_heading('Project Overview', level=2)
    doc.add_paragraph(
        'We are pleased to present this comprehensive proposal for developing a modern ERP system tailored '
        'specifically for your import-distribution business operations. This solution will streamline your entire '
        'workflow from procurement to payment collection, providing real-time visibility and control across all '
        'business functions.'
    )
    doc.add_paragraph(
        'Building on our expertise in enterprise software development, this ERP system will provide a robust, '
        'scalable platform that addresses the unique challenges of import-distribution operations while '
        'maintaining ease of use and reliability.'
    )

    # Project Scope
    doc.add_heading('Project Scope', level=2)
    doc.add_paragraph('Core Deliverables:')

    deliverables = [
        'Complete inventory management system with multi-warehouse support',
        'Purchase order and supplier management',
        'Sales invoicing with automated stock deduction',
        'Client management with credit limit enforcement',
        'Import documentation and landed cost tracking',
        'Recovery management with aging analysis',
        'Comprehensive reporting and analytics',
        'Role-based access control (5 user roles)',
        'Responsive web application (desktop, tablet, mobile)',
        'Complete audit trail and activity logging',
        'Excel export functionality across all reports'
    ]
    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')

    # Technology Stack
    doc.add_heading('Technology Stack', level=2)
    tech_para = doc.add_paragraph()
    tech_para.add_run('Frontend: ').bold = True
    tech_para.add_run('Modern web technologies for responsive, fast-loading interfaces\n')
    tech_para.add_run('Backend: ').bold = True
    tech_para.add_run('Robust server-side framework for scalability and performance\n')
    tech_para.add_run('Database: ').bold = True
    tech_para.add_run('Reliable database system for secure data management\n')
    tech_para.add_run('Deployment: ').bold = True
    tech_para.add_run('On-premise installation with complete source code access')

    add_page_break(doc)

    # ==================== PAGE 3: KEY FEATURES ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Key Features', level=1)

    # Business Management Features
    doc.add_heading('Business Management Features', level=2)

    business_features = [
        ('Multi-Warehouse Inventory Control', 'Real-time stock tracking across unlimited warehouses with bin location management, low stock alerts, and automated stock movements'),
        ('Import & Procurement Management', 'Comprehensive import documentation tracking with landed cost calculation (customs, taxes, shipping), supplier management, and purchase order workflows'),
        ('Batch/Lot Traceability', 'Complete product traceability from receipt to sale with expiry date tracking, batch-wise costing, and quality control'),
        ('Credit Limit Enforcement', 'Automated credit limit checks during sales with real-time balance updates, payment due tracking, and customer credit reports'),
        ('Weekly Recovery Schedule', 'Intelligent recovery planning with aging analysis, automated reminders, and collection performance tracking to reduce DSO by 30%'),
        ('Gate Pass System', 'Flexible approval workflows (auto/manual) for inventory dispatch with complete audit trail and delivery tracking')
    ]

    for title, desc in business_features:
        feature_para = doc.add_paragraph()
        feature_para.add_run(f'{title}: ').bold = True
        feature_para.add_run(desc)
        feature_para.style = 'List Bullet'

    # Customer Experience Features
    doc.add_heading('Customer Experience Features', level=2)

    customer_features = [
        ('Intelligent Product Catalog', 'Smart search with auto-suggestions, advanced filtering, and related product recommendations'),
        ('Flexible Sales Processing', 'Multi-step checkout with automated calculations, discount management, and multiple payment options'),
        ('Comprehensive User Accounts', 'Customer profiles with order history, address book management, and reorder functionality'),
        ('Real-time Order Tracking', 'Email notifications, order status updates, and return/exchange request handling'),
        ('Mobile-Optimized Interface', 'Responsive design works seamlessly on desktop, tablet, and mobile devices')
    ]

    for title, desc in customer_features:
        feature_para = doc.add_paragraph()
        feature_para.add_run(f'{title}: ').bold = True
        feature_para.add_run(desc)
        feature_para.style = 'List Bullet'

    # Administrative Control Center
    doc.add_heading('Administrative Control Center', level=2)

    admin_features = [
        ('Powerful Admin Dashboard', 'Real-time business metrics, sales analytics, and intuitive management interface'),
        ('Advanced Product Management', 'Bulk import/export, variant management, category organization, and product status control'),
        ('Order Processing System', 'Order workflow management, shipping integration, return processing, and customer communication'),
        ('Financial Tools', 'Tax management, discount creation, pricing controls, and comprehensive financial reporting'),
        ('Security & Access Control', 'Role-based permissions (Admin, Warehouse, Sales, Accountant, Recovery Agent) with complete activity logging')
    ]

    for title, desc in admin_features:
        feature_para = doc.add_paragraph()
        feature_para.add_run(f'{title}: ').bold = True
        feature_para.add_run(desc)
        feature_para.style = 'List Bullet'

    add_page_break(doc)

    # ==================== PAGE 4: DEVELOPMENT PHASES ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Development Phases', level=1)

    # Phase 1
    phase1 = doc.add_paragraph()
    phase1_label = phase1.add_run('Phase 1')
    phase1_label.font.size = Pt(10)
    phase1_label.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading('Foundation Setup', level=2)
    phase1_items = [
        'Framework installation and configuration',
        'Database setup and initial structure',
        'Basic user authentication implementation',
        'Project architecture and code structure'
    ]
    for item in phase1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Phase 2
    phase2 = doc.add_paragraph()
    phase2_label = phase2.add_run('Phase 2')
    phase2_label.font.size = Pt(10)
    phase2_label.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading('Core Functionality', level=2)
    phase2_items = [
        'Product catalog development',
        'User authentication and account management',
        'Inventory management system',
        'Purchase order and supplier management'
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Phase 3
    phase3 = doc.add_paragraph()
    phase3_label = phase3.add_run('Phase 3')
    phase3_label.font.size = Pt(10)
    phase3_label.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading('Advanced Features', level=2)
    phase3_items = [
        'Sales invoicing and client management',
        'Import documentation tracking',
        'Recovery schedule and aging reports',
        'Admin dashboard development'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Phase 4
    phase4 = doc.add_paragraph()
    phase4_label = phase4.add_run('Phase 4')
    phase4_label.font.size = Pt(10)
    phase4_label.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading('Optimization & Testing', level=2)
    phase4_items = [
        'Performance optimization',
        'Security implementation',
        'Cross-browser and device testing',
        'Quality assurance and bug fixes',
        'User acceptance testing and training'
    ]
    for item in phase4_items:
        doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # ==================== PAGE 5: HOSTING & INFRASTRUCTURE ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Hosting & Infrastructure', level=1)

    doc.add_heading('Infrastructure Requirements', level=2)
    doc.add_paragraph(
        'Your ERP system will require robust hosting infrastructure to ensure optimal performance, security, '
        'and reliability. We recommend on-premise or cloud-based solutions depending on your requirements.'
    )

    doc.add_heading('Recommended Setup', level=2)
    setup_items = [
        'Application Server: Dedicated server or cloud instance with sufficient resources',
        'Database Server: Managed database service with automated backups',
        'SSL Certificate: Security encryption for data protection',
        'Domain Setup: Professional domain configuration',
        'Backup System: Regular automated backups and recovery procedures'
    ]
    for item in setup_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Our Role', level=2)
    role_items = [
        'Setup Assistance: We will help configure and optimize your hosting environment',
        'Deployment Support: Seamless deployment from development to production',
        'Initial Configuration: Server optimization and security setup',
        'Documentation: Complete deployment and maintenance guides'
    ]
    for item in role_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        'Important Note: '
    )
    note_run.italic = True
    note_run.bold = True
    note_para.add_run(
        'Hosting and infrastructure costs are separate from our development fees and will be arranged '
        'based on your requirements and preferences.'
    ).italic = True

    doc.add_heading('Annual Maintenance Contract (AMC)', level=2)
    doc.add_paragraph("What's Included:")
    amc_items = [
        'Security Updates: Regular software and security patches',
        'Performance Monitoring: Ongoing system performance optimization',
        'Bug Fixes: Resolution of any technical issues that arise',
        'Content Updates: Minor content changes and feature tweaks',
        'Backup Management: Regular automated backups and recovery procedures',
        'Technical Support: Priority email support for technical queries',
        'Monthly Reports: Performance and analytics reporting'
    ]
    for item in amc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    amc_note = doc.add_paragraph()
    amc_note.add_run('Annual Maintenance Fee: ').bold = True
    amc_note.add_run('Available upon request after project completion')

    add_page_break(doc)

    # ==================== PAGE 6: PRICING ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Project Pricing', level=1)

    # Metadata
    pricing_meta = doc.add_paragraph()
    pricing_meta.add_run('Prepared for: ').bold = True
    pricing_meta.add_run('Hisham Traders\n')
    pricing_meta.add_run('Prepared by: ').bold = True
    pricing_meta.add_run('Sinnextech\n')
    pricing_meta.add_run('Date: ').bold = True
    pricing_meta.add_run('January 2025')

    doc.add_paragraph()

    # Pricing table
    pricing_table = doc.add_table(rows=6, cols=3)
    pricing_table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = pricing_table.rows[0].cells
    hdr_cells[0].text = 'Services'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Cost'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    # Row 1: Basic Tier
    row1 = pricing_table.rows[1].cells
    row1[0].text = 'Basic / MVP Tier'
    row1[1].text = 'Essential operations with single warehouse, basic inventory, sales invoicing, and 3 user roles'
    row1[2].text = '$1,150'

    # Row 2: Standard Tier
    row2 = pricing_table.rows[2].cells
    row2[0].text = 'Standard Tier\n(RECOMMENDED)'
    row2[1].text = 'Complete ERP with multi-warehouse, import tracking, batch/lot management, credit control, recovery schedule, and 5 user roles'
    row2[2].text = '$1,850'
    row2[0].paragraphs[0].runs[0].font.bold = True

    # Row 3: Premium Tier
    row3 = pricing_table.rows[3].cells
    row3[0].text = 'Premium Tier'
    row3[1].text = 'Enterprise features with barcode scanning, mobile apps, WhatsApp/SMS integration, AI analytics, and FBR integration'
    row3[2].text = '$3,400'

    # Row 4: Setup & Testing
    row4 = pricing_table.rows[4].cells
    row4[0].text = 'Setup & Testing'
    row4[1].text = 'Hosting configuration, deployment support, and comprehensive testing'
    row4[2].text = 'Included'

    # Row 5: Documentation
    row5 = pricing_table.rows[5].cells
    row5[0].text = 'Documentation'
    row5[1].text = 'Technical documentation and user guides'
    row5[2].text = 'Included'

    doc.add_paragraph()

    # Additional Services
    doc.add_heading('Additional Services', level=2)

    add_services_table = doc.add_table(rows=7, cols=3)
    add_services_table.style = 'Light List Accent 1'

    # Header
    add_hdr = add_services_table.rows[0].cells
    add_hdr[0].text = 'Service'
    add_hdr[1].text = 'Description'
    add_hdr[2].text = 'Price'

    for cell in add_hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    services = [
        ('Extended Support', '12-month priority support package', '$180/year'),
        ('Advanced Training', 'Role-specific training (per role)', '$90'),
        ('Custom Reports', 'Bespoke report development', '$55 each'),
        ('Data Migration', 'Legacy system data import', '$110 - 270'),
        ('API Integration', 'Third-party system integration', '$180 - 540'),
        ('Hosting Setup', 'Cloud/VPS server setup & configuration', '$145')
    ]

    for idx, (service, desc, price) in enumerate(services, 1):
        row_cells = add_services_table.rows[idx].cells
        row_cells[0].text = service
        row_cells[1].text = desc
        row_cells[2].text = price

    add_page_break(doc)

    # ==================== PAGE 7: PAYMENT SCHEDULE ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Payment Schedule', level=1)

    doc.add_paragraph(
        'Flexible payment terms designed to align with project milestones and deliverables:'
    )

    doc.add_paragraph()

    # Payment schedule table
    payment_table = doc.add_table(rows=4, cols=3)
    payment_table.style = 'Medium Grid 1 Accent 1'

    # Header
    pay_hdr = payment_table.rows[0].cells
    pay_hdr[0].text = 'Payment'
    pay_hdr[1].text = 'Amount'
    pay_hdr[2].text = 'Due Date'

    for cell in pay_hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    # Payment rows (using Standard tier as example)
    payment_data = [
        ('Advance Payment', '50% (Example: $925 for Standard)', 'Upon proposal acceptance'),
        ('Milestone Payment', '30% (Example: $555 for Standard)', 'Upon design completion + 50% development progress'),
        ('Final Payment', '20% (Example: $370 for Standard)', 'Upon project completion and launch')
    ]

    for idx, (payment, amount, due) in enumerate(payment_data, 1):
        row_cells = payment_table.rows[idx].cells
        row_cells[0].text = payment
        row_cells[1].text = amount
        row_cells[2].text = due

    doc.add_paragraph()

    doc.add_heading("What's Not Included", level=2)
    not_included = [
        'Major feature additions or modifications beyond agreed scope',
        'Third-party integration beyond original scope',
        'Hosting and infrastructure costs',
        'Marketing and advertising services',
        'Extensive design changes after approval'
    ]
    for item in not_included:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Additional Development & Integrations', level=2)
    doc.add_paragraph(
        'Any development work beyond the agreed scope will be quoted and charged separately based on '
        'complexity and time requirements. We will provide detailed estimates for any additional work before '
        'proceeding.'
    )

    add_page_break(doc)

    # ==================== PAGE 8: AGREEMENT & SIGNATURE ====================
    logo_header = doc.add_paragraph()
    logo_header_run = logo_header.add_run('SinnexTech')
    logo_header_run.font.size = Pt(14)
    logo_header_run.font.bold = True
    logo_header_run.font.color.rgb = RGBColor(102, 126, 234)

    doc.add_paragraph()

    doc.add_heading('Agreement & Next Steps', level=1)

    doc.add_paragraph(
        'By signing below, both parties agree to the terms, scope, and pricing outlined in this proposal. '
        'Upon receipt of the advance payment, we will commence work immediately.'
    )

    doc.add_paragraph()

    doc.add_heading('Project Timeline: 6-10 weeks from advance payment receipt', level=2)

    doc.add_paragraph()

    doc.add_heading('What Happens Next:', level=2)
    next_steps = [
        'Proposal Acceptance: Sign this document and submit advance payment',
        'Project Commencement: We begin work upon payment confirmation',
        'Regular Updates: Weekly progress reports and milestone reviews',
        'Quality Delivery: Comprehensive testing and launch support'
    ]

    for idx, step in enumerate(next_steps, 1):
        step_para = doc.add_paragraph()
        step_para.add_run(f'{idx}. {step}')

    doc.add_paragraph()
    doc.add_paragraph()

    # Proposal Acceptance section
    doc.add_heading('Proposal Acceptance', level=2)

    doc.add_paragraph()

    # Two column layout simulation
    acceptance_table = doc.add_table(rows=1, cols=2)
    acceptance_table.autofit = False

    # Left column - Client
    left_cell = acceptance_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.add_run('Client Acceptance:').bold = True
    left_para.add_run('\n\n')
    left_para.add_run('Date: ________________\n\n')
    left_para.add_run('Client Name: ________________\n\n')
    left_para.add_run('Company: Hisham Traders\n\n')
    left_para.add_run('Signature: ________________')

    # Right column - Sinnextech
    right_cell = acceptance_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.add_run('Sincerely,').bold = True
    right_para.add_run('\n\n\n\n')
    right_para.add_run('Haisam Shoaib\n').bold = True
    right_para.add_run('CEO at Sinnextech')

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing note
    closing = doc.add_paragraph()
    closing.add_run(
        "We're excited to partner with you and bring your ERP vision to life. This platform will serve "
        "as a powerful management tool that streamlines your operations and drives business growth."
    ).italic = True

    doc.add_paragraph()

    # Validity note
    validity_box = doc.add_paragraph()
    validity_box.add_run(
        'This proposal is valid for 30 days from the date of submission. All prices are quoted in USD and '
        'exclude applicable taxes.'
    )
    # Add light background color effect with paragraph spacing
    validity_box.paragraph_format.space_before = Pt(12)
    validity_box.paragraph_format.space_after = Pt(12)
    validity_box.paragraph_format.left_indent = Pt(12)
    validity_box.paragraph_format.right_indent = Pt(12)

    # Add footer to all sections
    add_footer(doc)

    # Save document
    doc.save('e:\\pProjects\\hishamtraders\\demo\\Hisham_Traders_ERP_Proposal.docx')
    print("[SUCCESS] Proposal document created successfully with Sinnextech branding!")
    print("Location: demo/Hisham_Traders_ERP_Proposal.docx")

if __name__ == '__main__':
    try:
        create_proposal()
    except ImportError:
        print("Error: python-docx library not found")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
